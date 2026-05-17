#!/usr/bin/env python3
"""Construye Respawn_SDD_v2_final.docx replicando el parcial + correcciones
globales + inyección de secciones nuevas del JSON generado en PASO 1.

Salida: Respawn_SDD_v2_final.docx (raíz del proyecto).
"""

import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
SRC = Path(r"C:\Users\nicol\Downloads\Respawn_Entrega_Parcial.docx")
JSON_NUEVAS = ROOT / ".tmp-nuevas-secciones.json"
OUT = ROOT / "Respawn_SDD_v2_final.docx"

# ─── Correcciones globales ──────────────────────────────────────────────────
CORRECCIONES = [
    ("Next.js 15", "Next.js 16.2.3"),
    ("Tailwind CSS v4", "Tailwind CSS v3.4"),
    ("cost factor mínimo de 12", "cost factor de 10"),
]

def aplicar_correcciones(texto: str) -> str:
    if not texto:
        return texto
    out = texto
    for old, new in CORRECCIONES:
        out = out.replace(old, new)
    return out


# ─── Setup formato CEAC ─────────────────────────────────────────────────────
def configurar_documento(doc: Document) -> None:
    """Aplica formato CEAC: A4, márgenes, Calibri 12, interlineado 1.5."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Estilo Normal
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(6)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ─── Helpers de bloques tipados (sección NUEVA) ─────────────────────────────
def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x39, 0x7D)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)

def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x4D, 0x9B)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)

def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(12.5)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)

def add_h4(doc: Document, text: str):
    """Estilo de label de figura/tabla."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_p(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(aplicar_correcciones(text))
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_p_italic(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.italic = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

def add_p_small_center(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(9.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)

def add_li(doc: Document, text: str):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(aplicar_correcciones(text))
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_li_apa(doc: Document, text: str):
    """Referencia APA: sangría francesa (primera línea sin sangría, resto sangrado)."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_code(doc: Document, text: str):
    """Bloque monoespaciado para árboles de carpetas o snippets."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    for line in text.split('\n'):
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        p.add_run('\n')

def add_review(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(f"[{text}]")
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


RENDERERS = {
    'h1': add_h1,
    'h2': add_h2,
    'h3': add_h3,
    'h4': add_h4,
    'p': add_p,
    'p_italic': add_p_italic,
    'p_small_center': add_p_small_center,
    'li': add_li,
    'li_num': add_li,  # numerado no necesario, reutilizamos bullet
    'li_apa': add_li_apa,
    'code': add_code,
    'review': add_review,
}

def render_blocks(doc: Document, blocks: list) -> None:
    for b in blocks:
        kind = b['kind']
        text = b.get('text', '')
        fn = RENDERERS.get(kind, add_p)
        fn(doc, text)


# ─── Copia del parcial preservando estructura y aplicando correcciones ──────
def copy_paragraph(src_p, dst_doc: Document):
    """Crea un nuevo párrafo en dst_doc replicando estilo y texto (corregido)."""
    text = src_p.text
    if not text.strip() and not src_p.style.name.startswith('Heading'):
        # Párrafo vacío: añadimos un break suave para preservar separación
        dst_doc.add_paragraph()
        return

    style_name = src_p.style.name if src_p.style else 'Normal'
    text_corr = aplicar_correcciones(text)

    if style_name == 'Title':
        # Título del documento
        p = dst_doc.add_paragraph()
        r = p.add_run(text_corr)
        r.font.name = 'Calibri'
        r.font.size = Pt(24)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1F, 0x39, 0x7D)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(18)
        return

    if style_name == 'Heading 1':
        add_h1(dst_doc, text_corr)
        return
    if style_name == 'Heading 2':
        add_h2(dst_doc, text_corr)
        return
    if style_name == 'Heading 3':
        add_h3(dst_doc, text_corr)
        return
    if style_name == 'Heading 4':
        add_h4(dst_doc, text_corr)
        return
    if style_name == 'List Paragraph':
        add_li(dst_doc, text_corr)
        return

    # Párrafo normal
    add_p(dst_doc, text_corr)


def copy_table(src_t, dst_doc: Document):
    """Replica una tabla aplicando correcciones a cada celda."""
    n_rows = len(src_t.rows)
    n_cols = len(src_t.columns)
    if n_rows == 0 or n_cols == 0:
        return
    t = dst_doc.add_table(rows=n_rows, cols=n_cols)
    t.style = 'Table Grid'

    for r_i, row in enumerate(src_t.rows):
        for c_i in range(n_cols):
            if c_i < len(row.cells):
                src_cell = row.cells[c_i]
                cell_text = aplicar_correcciones(src_cell.text)
                dst_cell = t.rows[r_i].cells[c_i]
                # Limpiar contenido por defecto
                dst_cell.text = ''
                p = dst_cell.paragraphs[0]
                # Cabecera (primera fila) en negrita y sombreado gris
                if r_i == 0:
                    r = p.add_run(cell_text)
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
                    r.font.bold = True
                    # Sombreado de cabecera
                    tcPr = dst_cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'E7E6E6')
                    tcPr.append(shd)
                else:
                    r = p.add_run(cell_text)
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)
    # Espacio después de tabla
    dst_doc.add_paragraph()


# ─── Punto de inserción de las secciones nuevas ─────────────────────────────
# El parcial termina con la sección DISEÑO. Insertamos:
#   1. "diseno_completar" justo después del último párrafo del parcial
#      (que es la nota sobre la vista post_scores).
#   2. Implementación, Conclusiones, Innovación, Bibliografía, Anexos
#      a continuación.
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print(f"Leyendo parcial: {SRC}")
    src = Document(str(SRC))

    print("Creando documento destino...")
    dst = Document()
    configurar_documento(dst)

    # ── Copiar contenido del parcial en orden ──
    body = src.element.body
    paragraphs = list(src.paragraphs)
    tables = list(src.tables)
    p_idx = 0
    t_idx = 0
    copiados_p = 0
    copiados_t = 0

    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            if p_idx < len(paragraphs):
                copy_paragraph(paragraphs[p_idx], dst)
                copiados_p += 1
            p_idx += 1
        elif tag == 'tbl':
            if t_idx < len(tables):
                copy_table(tables[t_idx], dst)
                copiados_t += 1
            t_idx += 1
        elif tag == 'sectPr':
            # Saltos de sección los ignoramos (formato CEAC se aplica al doc entero)
            pass

    print(f"Copiados del parcial: {copiados_p} parrafos, {copiados_t} tablas")

    # ── Insertar secciones nuevas ──
    print("Cargando secciones nuevas...")
    nuevas = json.loads(JSON_NUEVAS.read_text(encoding='utf-8'))

    # Orden de inserción
    orden = [
        ('diseno_completar', 'Completar DISEÑO con diagrama secuencia post + Prototipo'),
        ('implementacion',   'Sección IMPLEMENTACIÓN'),
        ('conclusiones',     'Sección CONCLUSIONES'),
        ('innovacion',       'Sección INNOVACIÓN Y PROSPECTIVA'),
        ('bibliografia',     'Sección BIBLIOGRAFÍA'),
        ('anexos',           'Sección ANEXOS'),
    ]
    for clave, descripcion in orden:
        blocks = nuevas[clave]
        print(f"  -> {clave}: {len(blocks)} bloques")
        render_blocks(dst, blocks)

    # ── Guardar ──
    print(f"Guardando: {OUT}")
    dst.save(str(OUT))
    print(f"\nOK -> {OUT}")
    print(f"Tamaño: {OUT.stat().st_size} bytes")

if __name__ == '__main__':
    main()
