#!/usr/bin/env python3
"""Valida que el docx generado abre correctamente, cuenta parrafos/tablas,
estima paginas, comprueba correcciones aplicadas y lista placeholders."""

from pathlib import Path
from docx import Document

OUT = Path(__file__).parent.parent / "Respawn_SDD_v2_final.docx"

print(f"Abriendo: {OUT}")
doc = Document(str(OUT))

# Conteos
p_count = len(doc.paragraphs)
t_count = len(doc.tables)
words = 0
heading_counts = {}
for p in doc.paragraphs:
    if p.text.strip():
        words += len(p.text.split())
    style = p.style.name if p.style else 'Normal'
    if style.startswith('Heading') or style in ('Title', 'List Bullet'):
        heading_counts[style] = heading_counts.get(style, 0) + 1
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            words += len(cell.text.split())

# Estimación páginas (250 palabras/página A4 Calibri 12 interlineado 1.5)
pages_estimate = max(1, round(words / 260))

# Verificación de correcciones aplicadas
all_text = '\n'.join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            all_text += '\n' + cell.text

issues = []
for old in ["Next.js 15", "Tailwind CSS v4", "cost factor mínimo de 12"]:
    if old in all_text:
        # Excepción: "Next.js 15" puede aparecer dentro de "Next.js 15."
        # como parte de error legítimo, pero el match exacto debería ser 0
        ocurrencias = all_text.count(old)
        issues.append(f"  AVISO: queda(n) {ocurrencias} mencion(es) de '{old}'")

# Verificación de nuevas referencias correctas
for new in ["Next.js 16.2.3", "Tailwind CSS v3.4", "cost factor de 10"]:
    if new not in all_text:
        issues.append(f"  AVISO: '{new}' NO encontrado (correccion no aplicada)")

# Placeholders REVISAR
import re
revisar = re.findall(r'\[REVISAR:[^\]]+\]', all_text)

# Verificar secciones nuevas presentes
secciones_nuevas = ['IMPLEMENTACIÓN', 'CONCLUSIONES', 'INNOVACIÓN Y PROSPECTIVA', 'BIBLIOGRAFÍA', 'ANEXOS']
missing = [s for s in secciones_nuevas if s not in all_text]

print("\n=== ESTADÍSTICAS ===")
print(f"Parrafos: {p_count}")
print(f"Tablas: {t_count}")
print(f"Palabras totales: {words}")
print(f"Estimacion paginas: ~{pages_estimate} paginas A4")
print(f"\nDistribucion de estilos:")
for s, n in sorted(heading_counts.items()):
    print(f"  {s}: {n}")

print(f"\n=== CORRECCIONES GLOBALES ===")
if not issues:
    print("OK Todas las correcciones aplicadas correctamente")
else:
    for i in issues:
        print(i)

print(f"\n=== SECCIONES NUEVAS ===")
for s in secciones_nuevas:
    status = "OK" if s not in missing else "FALTA"
    print(f"  [{status}] {s}")

print(f"\n=== PLACEHOLDERS [REVISAR] ({len(revisar)}) ===")
for r in revisar:
    print(f"  {r}")

# Resumen final
print(f"\n=== ARCHIVO ===")
print(f"Ruta: {OUT}")
print(f"Tamano: {OUT.stat().st_size:,} bytes")
print(f"\nValidacion: {'PASS' if not issues and not missing else 'CON AVISOS'}")
