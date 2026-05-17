#!/usr/bin/env python3
"""Vuelca la estructura del .docx parcial a un .txt legible para que Claude vea
todo el contenido sin perder formato semántico (títulos, tablas, listas)."""

import sys
from docx import Document

SRC = r"C:\Users\nicol\Downloads\Respawn_Entrega_Parcial.docx"
OUT = r"C:\Users\nicol\Desktop\Forogaming\.tmp-parcial-dump.txt"

doc = Document(SRC)

lines = []
lines.append(f"=== DOCX: {SRC} ===\n")
lines.append(f"Total parrafos: {len(doc.paragraphs)}")
lines.append(f"Total tablas: {len(doc.tables)}")
lines.append(f"Total secciones: {len(doc.sections)}\n")

# Estilos del documento (para conocer la jerarquía)
for sec_i, sec in enumerate(doc.sections):
    lines.append(f"--- Section {sec_i}: margins L={sec.left_margin} R={sec.right_margin} T={sec.top_margin} B={sec.bottom_margin}")
lines.append("")

# Iteramos en orden de aparición preservando body
body = doc.element.body
para_count = 0
table_count = 0

def safe(text):
    return text if text else ""

for child in body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        if para_count < len(doc.paragraphs):
            p = doc.paragraphs[para_count]
        else:
            p = None
        para_count += 1
        if p is None:
            continue
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()
        if not text and not style.startswith('Heading'):
            continue
        # Indicar estilo de cabecera explícitamente
        if style.startswith('Heading'):
            lines.append(f"[{style}] {text}")
        elif style == 'Title':
            lines.append(f"[TITLE] {text}")
        elif style == 'Subtitle':
            lines.append(f"[SUBTITLE] {text}")
        elif style == 'List Paragraph':
            lines.append(f"  - {text}")
        else:
            lines.append(text)
    elif tag == 'tbl':
        if table_count < len(doc.tables):
            t = doc.tables[table_count]
        else:
            t = None
        table_count += 1
        if t is None:
            continue
        lines.append(f"\n[TABLA #{table_count}] {len(t.rows)} filas x {len(t.columns)} cols")
        for r_i, row in enumerate(t.rows):
            row_cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
            lines.append(f"  R{r_i}: {' || '.join(row_cells)}")
        lines.append("")

with open(OUT, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print(f"OK -> {OUT}  ({len(lines)} lineas, {para_count} parrafos, {table_count} tablas)")
